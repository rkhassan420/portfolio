from django.contrib.auth import authenticate
from django.contrib.auth.models import User
from django.shortcuts import get_object_or_404, redirect
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework_simplejwt.tokens import RefreshToken
from .models import HomeInfo, AboutInfo, FooterInfo, ProjectsInfo, OTPVerification, PortfolioLink, PortfolioVisit, \
    ProjectDetail
from .serializer import HomeSerializer, AboutSerializer, FooterSerializer, ProjectsSerializer, PortfolioLinkSerializer, \
    ProjectDetailSerializer, SettingsSerializer, CVPreferenceSerializer
from .utils import generate_otp, send_otp_email
from django.db.models import Count
from django.db.models.functions import TruncDate
from django.utils import timezone
from datetime import timedelta
from rest_framework.permissions import IsAuthenticated
import json
from django.contrib.auth import get_user_model
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle, KeepTogether
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Indenter
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    HRFlowable, KeepTogether, Table, TableStyle
)
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from rest_framework.views import APIView
from rest_framework import status
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse
import io

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from rest_framework.permissions import AllowAny
from django.shortcuts import get_object_or_404
from .models import CVPreference
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
import io

from .models import HomeInfo, AboutInfo, FooterInfo, ProjectsInfo, CVPreference, ALL_SECTIONS





# ─────────────────────────────────────────────
#  OTP
# ─────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([AllowAny])
def SendOTPView(request):
    email = request.data.get('email', '').strip()

    if not email:
        return Response({'error': 'Email is required'}, status=400)
    if '@' not in email or '.' not in email:
        return Response({'error': 'Invalid email format'}, status=400)

    otp = generate_otp()
    OTPVerification.objects.filter(email=email).delete()
    OTPVerification.objects.create(email=email, otp=otp)

    try:
        send_otp_email(email, otp)
        return Response({'message': 'OTP sent successfully'}, status=200)
    except Exception as e:
        return Response({'error': f'Failed to send OTP: {str(e)}'}, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
def VerifyOTPView(request):
    email = request.data.get('email', '').strip()
    otp = request.data.get('otp', '').strip()

    if not email or not otp:
        return Response({'error': 'Email and OTP are required'}, status=400)

    try:
        record = OTPVerification.objects.filter(
            email=email, otp=otp, is_verified=False
        ).latest('created_at')

        if record.is_expired():
            return Response({'error': 'OTP has expired. Please request a new one.'}, status=400)

        record.is_verified = True
        record.save()
        return Response({'message': 'OTP verified successfully'}, status=200)

    except OTPVerification.DoesNotExist:
        return Response({'error': 'Invalid OTP. Please try again.'}, status=400)


# ─────────────────────────────────────────────
#  AUTH — JWT (no CSRF needed)
# ─────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([AllowAny])
def JWTRegisterView(request):
    username = request.data.get('username', '').strip()
    email = request.data.get('email', '').strip()
    password = request.data.get('password', '')

    if not all([username, email, password]):
        return Response({'error': 'All fields are required'}, status=400)
    if len(username) < 3:
        return Response({'error': 'Username must be at least 3 characters'}, status=400)
    if not username.replace('_', '').isalnum():
        return Response({'error': 'Username can only contain letters, numbers, underscores'}, status=400)
    if len(password) < 6:
        return Response({'error': 'Password must be at least 6 characters'}, status=400)
    if User.objects.filter(username=username).exists():
        return Response({'error': 'Username already exists'}, status=400)
    if User.objects.filter(email=email).exists():
        return Response({'error': 'Email already registered'}, status=400)

    user = User(username=username, email=email)
    user.set_password(password)
    user.save()

    refresh = RefreshToken.for_user(user)
    return Response({
        'message': 'Account created successfully',
        'username': user.username,
        'access': str(refresh.access_token),
        'refresh': str(refresh),
    }, status=201)


@api_view(['POST'])
@permission_classes([AllowAny])
def JWTLoginView(request):
    username = request.data.get('username', '').strip()
    password = request.data.get('password', '')

    if not all([username, password]):
        return Response({'error': 'Username and password are required'}, status=400)

    user = authenticate(request, username=username, password=password)
    if user is not None:
        refresh = RefreshToken.for_user(user)
        return Response({
            'message': 'Login successful',
            'username': user.username,
            'access': str(refresh.access_token),
            'refresh': str(refresh),
        }, status=200)
    else:
        return Response({'error': 'Invalid username or password'}, status=401)


@api_view(['POST'])
@permission_classes([AllowAny])
def RefreshTokenView(request):
    refresh_token = request.data.get('refresh')
    if not refresh_token:
        return Response({'error': 'Refresh token required'}, status=400)
    try:
        refresh = RefreshToken(refresh_token)
        return Response({'access': str(refresh.access_token)}, status=200)
    except Exception:
        return Response({'error': 'Invalid or expired refresh token'}, status=401)


# ─────────────────────────────────────────────
#  HOME INFO
# ─────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([AllowAny])
def get_home_info(request):
    username = request.query_params.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)
    try:
        obj = HomeInfo.objects.get(username=username)
        return Response(HomeSerializer(obj).data)
    except HomeInfo.DoesNotExist:
        return Response({"error": "Not found"}, status=404)


@api_view(['POST'])
@permission_classes([AllowAny])
def add_home_info(request):
    username = request.data.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)

    try:
        obj = HomeInfo.objects.get(username=username)
        data = request.data.copy()
        if not data.get('image'): data['image'] = obj.image
        if not data.get('cv'):    data['cv'] = obj.cv
        serializer = HomeSerializer(obj, data=data, partial=True)
    except HomeInfo.DoesNotExist:
        serializer = HomeSerializer(data=request.data)

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=200)
    return Response(serializer.errors, status=400)


# ─────────────────────────────────────────────
#  ABOUT INFO
# ─────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([AllowAny])
def get_about_info(request):
    username = request.query_params.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)
    try:
        obj = AboutInfo.objects.get(username=username)
        return Response(AboutSerializer(obj).data)
    except AboutInfo.DoesNotExist:
        return Response({"error": "Not found"}, status=404)


@api_view(['POST'])
@permission_classes([AllowAny])
def add_about_info(request):
    username = request.data.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)

    try:
        obj = AboutInfo.objects.get(username=username)
        data = request.data.copy()
        if not data.get('image'): data['image'] = obj.image
        serializer = AboutSerializer(obj, data=data, partial=True)
    except AboutInfo.DoesNotExist:
        serializer = AboutSerializer(data=request.data)

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=200)
    return Response(serializer.errors, status=400)


# ─────────────────────────────────────────────
#  FOOTER INFO
# ─────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([AllowAny])
def get_footer_info(request):
    username = request.query_params.get('username')
    if not username:
        return Response({'error': 'Username is required'}, status=400)
    try:
        objs = FooterInfo.objects.filter(username=username)
        return Response(FooterSerializer(objs, many=True).data, status=200)
    except Exception as e:
        return Response({'error': str(e)}, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
def add_footer_info(request):
    username = request.data.get('username')

    if not username:
        return Response({"error": "Username is required"}, status=400)

    data = {
        key: value
        for key, value in request.data.items()
        if value != ''
    }

    try:
        obj = FooterInfo.objects.get(username=username)
        serializer = FooterSerializer(obj, data=data, partial=True)

    except FooterInfo.DoesNotExist:
        serializer = FooterSerializer(data=data)

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=200)

    return Response(serializer.errors, status=400)


# ─────────────────────────────────────────────
#  PROJECTS INFO
# ─────────────────────────────────────────────

@api_view(['GET'])
@permission_classes([AllowAny])
def get_projects_info(request):
    username = request.query_params.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)

    objs = ProjectsInfo.objects.filter(username=username).order_by('-id')
    return Response({'data': ProjectsSerializer(objs, many=True).data, 'count': objs.count()})


@api_view(['POST'])
@permission_classes([AllowAny])
def add_projects_info(request):
    username = request.data.get('username')
    if not username:
        return Response({"error": "Username is required"}, status=400)

    serializer = ProjectsSerializer(data=request.data.copy())
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['DELETE'])
@permission_classes([AllowAny])
def delete_projects_info(request, pk):
    try:
        ProjectsInfo.objects.get(pk=pk).delete()
        return Response(status=204)
    except ProjectsInfo.DoesNotExist:
        return Response({'error': 'Project not found'}, status=404)


@api_view(['PUT'])
@permission_classes([AllowAny])
def edit_projects_info(request, pk):
    try:
        project = ProjectsInfo.objects.get(pk=pk)
    except ProjectsInfo.DoesNotExist:
        return Response({'error': 'Project not found'}, status=404)

    project.p_name = request.data.get('p_name', project.p_name)
    project.p_skills = request.data.get('p_skills', project.p_skills)
    project.p_url = request.data.get('p_url', project.p_url)
    image_url = request.data.get('image')
    if image_url:
        project.image = image_url
    project.save()
    return Response(ProjectsSerializer(project).data, status=200)


# ─────────────────────────────────────────────
#  ACCOUNT
# ─────────────────────────────────────────────

@api_view(['DELETE'])
@permission_classes([AllowAny])
def delete_user_account(request, username):
    try:
        User.objects.get(username=username).delete()
        return Response({'message': 'User deleted successfully'}, status=200)
    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=404)


# ─────────────────────────────────────────────
#  ANALYTICS
# ─────────────────────────────────────────────


@api_view(['POST'])
@permission_classes([AllowAny])
def track_visit(request):
    """Called automatically when someone opens a portfolio page."""
    username = request.data.get('username', '').strip()
    if not username:
        return Response({'error': 'Username required'}, status=400)

    PortfolioVisit.objects.create(
        username=username,
        country=request.data.get('country', ''),
        device=request.data.get('device', ''),
        browser=request.data.get('browser', ''),
        section=request.data.get('section', 'home'),
    )
    return Response({'message': 'Visit tracked'}, status=201)


@api_view(['GET'])
@permission_classes([AllowAny])
def get_analytics(request, username):  # username comes from URL
    username = username.strip()
    if not username:
        return Response({'error': 'Username required'}, status=400)

    visits = PortfolioVisit.objects.filter(username=username)

    total_visits = visits.count()

    seven_days_ago = timezone.now() - timedelta(days=6)
    daily = (
        visits.filter(visited_at__gte=seven_days_ago)
        .annotate(date=TruncDate('visited_at'))
        .values('date')
        .annotate(count=Count('id'))
        .order_by('date')
    )

    # Fill in missing days with 0
    daily_map = {str(d['date']): d['count'] for d in daily}
    daily_data = []
    for i in range(7):
        day = (timezone.now() - timedelta(days=6 - i)).date()
        daily_data.append({
            'date': str(day),
            'label': day.strftime('%a'),
            'count': daily_map.get(str(day), 0),
        })

    # ── Today vs Yesterday ──
    today = timezone.now().date()
    yesterday = today - timedelta(days=1)
    today_count = visits.filter(visited_at__date=today).count()
    yesterday_count = visits.filter(visited_at__date=yesterday).count()

    countries = (
        visits.exclude(country='')
        .values('country')
        .annotate(count=Count('id'))
        .order_by('-count')[:5]
    )

    # ── Device breakdown ──
    devices = (
        visits.exclude(device='')
        .values('device')
        .annotate(count=Count('id'))
        .order_by('-count')
    )

    # ── Browser breakdown ──
    browsers = (
        visits.exclude(browser='')
        .values('browser')
        .annotate(count=Count('id'))
        .order_by('-count')[:5]
    )

    # ── This month vs last month ──
    this_month = timezone.now().replace(day=1)
    last_month = (this_month - timedelta(days=1)).replace(day=1)
    this_month_count = visits.filter(visited_at__gte=this_month).count()
    last_month_count = visits.filter(
        visited_at__gte=last_month,
        visited_at__lt=this_month
    ).count()

    return Response({
        'total_visits': total_visits,
        'today_visits': today_count,
        'yesterday_visits': yesterday_count,
        'this_month_visits': this_month_count,
        'last_month_visits': last_month_count,
        'daily_data': daily_data,
        'top_countries': list(countries),
        'devices': list(devices),
        'browsers': list(browsers),
    }, status=200)


# ─────────────────────────────────────────────
#  OAUTH — Google & GitHub
# ─────────────────────────────────────────────

import requests as http_requests
from django.contrib.auth.models import User
from rest_framework_simplejwt.tokens import RefreshToken


def _jwt_for_user(user):
    """Helper — returns JWT tokens dict for a user."""
    refresh = RefreshToken.for_user(user)
    return {
        'username': user.username,
        'email': user.email,
        'access': str(refresh.access_token),
        'refresh': str(refresh),
    }


@api_view(['POST'])
@permission_classes([AllowAny])
def GoogleOAuthView(request):
    """
    Receives a Google access_token from frontend.
    Verifies it with Google, finds or creates the user, returns JWT.
    """
    access_token = request.data.get('access_token', '').strip()
    if not access_token:
        return Response({'error': 'access_token is required'}, status=400)

    # Verify token with Google
    try:
        google_resp = http_requests.get(
            'https://www.googleapis.com/oauth2/v3/userinfo',
            headers={'Authorization': f'Bearer {access_token}'},
            timeout=10,
        )
        if google_resp.status_code != 200:
            return Response({'error': 'Invalid Google token'}, status=401)

        google_data = google_resp.json()
    except Exception:
        return Response({'error': 'Failed to verify Google token'}, status=500)

    email = google_data.get('email', '').strip()
    name = google_data.get('name', '')
    google_id = google_data.get('sub', '')  # unique Google user ID

    if not email:
        return Response({'error': 'Could not retrieve email from Google'}, status=400)

    # Find or create user
    # Username = first part of email, made unique if needed
    base_username = email.split('@')[0].replace('.', '_').lower()
    user = User.objects.filter(email=email).first()

    if not user:
        username = base_username
        # Make username unique
        counter = 1
        while User.objects.filter(username=username).exists():
            username = f"{base_username}{counter}"
            counter += 1

        user = User.objects.create(
            username=username,
            email=email,
            first_name=name.split()[0] if name else '',
            last_name=' '.join(name.split()[1:]) if name and len(name.split()) > 1 else '',
        )
        user.set_unusable_password()  # OAuth users don't have passwords
        user.save()

    return Response({'message': 'Google login successful', **_jwt_for_user(user)}, status=200)


@api_view(['POST'])
@permission_classes([AllowAny])
def GitHubOAuthView(request):
    """
    Receives a GitHub code from frontend (after OAuth redirect).
    Exchanges code for access_token, fetches user info, returns JWT.
    """
    code = request.data.get('code', '').strip()
    if not code:
        return Response({'error': 'code is required'}, status=400)

    client_id = request.data.get('client_id', '').strip()
    client_secret = request.data.get('client_secret', '').strip()

    if not client_id or not client_secret:
        return Response({'error': 'client_id and client_secret required'}, status=400)

    # Exchange code for access token
    try:
        token_resp = http_requests.post(
            'https://github.com/login/oauth/access_token',
            data={
                'client_id': client_id,
                'client_secret': client_secret,
                'code': code,
            },
            headers={'Accept': 'application/json'},
            timeout=10,
        )
        token_data = token_resp.json()
        access_token = token_data.get('access_token')

        if not access_token:
            return Response({'error': 'Failed to get GitHub access token'}, status=401)

    except Exception:
        return Response({'error': 'Failed to contact GitHub'}, status=500)

    # Get user info
    try:
        user_resp = http_requests.get(
            'https://api.github.com/user',
            headers={
                'Authorization': f'Bearer {access_token}',
                'Accept': 'application/vnd.github.v3+json',
            },
            timeout=10,
        )
        github_data = user_resp.json()

        # GitHub may hide email — fetch primary email separately
        email = github_data.get('email')
        if not email:
            emails_resp = http_requests.get(
                'https://api.github.com/user/emails',
                headers={
                    'Authorization': f'Bearer {access_token}',
                    'Accept': 'application/vnd.github.v3+json',
                },
                timeout=10,
            )
            emails = emails_resp.json()
            primary = next((e for e in emails if e.get('primary') and e.get('verified')), None)
            email = primary['email'] if primary else None

    except Exception:
        return Response({'error': 'Failed to fetch GitHub user info'}, status=500)

    github_login = github_data.get('login', '')  # GitHub username
    name = github_data.get('name', '') or github_login

    if not email and not github_login:
        return Response({'error': 'Could not retrieve user info from GitHub'}, status=400)

    # Find or create user — prefer matching by email, fallback to github username
    user = None
    if email:
        user = User.objects.filter(email=email).first()

    if not user:
        user = User.objects.filter(username=github_login).first()

    if not user:
        # Create new user
        base_username = github_login or (email.split('@')[0] if email else 'user')
        username = base_username
        counter = 1
        while User.objects.filter(username=username).exists():
            username = f"{base_username}{counter}"
            counter += 1

        user = User.objects.create(
            username=username,
            email=email or '',
            first_name=name.split()[0] if name else '',
            last_name=' '.join(name.split()[1:]) if name and len(name.split()) > 1 else '',
        )
        user.set_unusable_password()
        user.save()

    return Response({'message': 'GitHub login successful', **_jwt_for_user(user)}, status=200)


# ─────────────────────────────────────────────
#  FORGOT PASSWORD
# ─────────────────────────────────────────────

@api_view(['POST'])
@permission_classes([AllowAny])
def ForgotPasswordSendOTP(request):
    """Send OTP to email for password reset."""
    email = request.data.get('email', '').strip()

    if not email:
        return Response({'error': 'Email is required'}, status=400)

    # Check user exists
    if not User.objects.filter(email=email).exists():
        # Don't reveal if email exists — security best practice
        return Response({'message': 'If this email exists, an OTP has been sent.'}, status=200)

    otp = generate_otp()
    OTPVerification.objects.filter(email=email).delete()
    OTPVerification.objects.create(email=email, otp=otp)

    try:
        from django.core.mail import send_mail
        from django.conf import settings
        send_mail(
            subject='ShowCraft — Password Reset Code',
            message=f"""
Hello,

You requested a password reset for your ShowCraft account.

Your reset code is:

  {otp}

This code is valid for 5 minutes.

If you did not request this, please ignore this email.

— The ShowCraft Team
""",
            from_email=settings.EMAIL_HOST_USER,
            recipient_list=[email],
            fail_silently=False,
        )
        return Response({'message': 'If this email exists, an OTP has been sent.'}, status=200)
    except Exception as e:
        return Response({'error': f'Failed to send email: {str(e)}'}, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
def ForgotPasswordVerifyOTP(request):
    """Verify OTP for password reset."""
    email = request.data.get('email', '').strip()
    otp = request.data.get('otp', '').strip()

    if not email or not otp:
        return Response({'error': 'Email and OTP are required'}, status=400)

    try:
        record = OTPVerification.objects.filter(
            email=email, otp=otp, is_verified=False
        ).latest('created_at')

        if record.is_expired():
            return Response({'error': 'OTP has expired. Please request a new one.'}, status=400)

        # Mark as verified but don't delete yet — needed for reset step
        record.is_verified = True
        record.save()
        return Response({'message': 'OTP verified. You may now reset your password.'}, status=200)

    except OTPVerification.DoesNotExist:
        return Response({'error': 'Invalid OTP. Please try again.'}, status=400)


@api_view(['POST'])
@permission_classes([AllowAny])
def ForgotPasswordReset(request):
    """Reset password after OTP verified."""
    email = request.data.get('email', '').strip()
    otp = request.data.get('otp', '').strip()
    new_password = request.data.get('new_password', '')

    if not all([email, otp, new_password]):
        return Response({'error': 'Email, OTP, and new password are required'}, status=400)

    if len(new_password) < 6:
        return Response({'error': 'Password must be at least 6 characters'}, status=400)

    # Re-verify the OTP is verified
    verified = OTPVerification.objects.filter(
        email=email, otp=otp, is_verified=True
    ).exists()

    if not verified:
        return Response({'error': 'OTP not verified. Please restart the process.'}, status=400)

    try:
        user = User.objects.get(email=email)
        user.set_password(new_password)
        user.save()

        # Clean up OTP
        OTPVerification.objects.filter(email=email).delete()

        return Response({'message': 'Password reset successfully. You can now log in.'}, status=200)

    except User.DoesNotExist:
        return Response({'error': 'User not found'}, status=404)


class CheckSlugView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        slug = request.query_params.get('slug', '').strip()
        if not slug:
            return Response({'error': 'Slug is required'}, status=400)
        if len(slug) < 3:
            return Response({'available': False, 'error': 'Min 3 characters'}, status=200)
        taken = PortfolioLink.objects.filter(slug=slug).exclude(user=request.user).exists()
        return Response({'available': not taken})


class PortfolioLinkView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        try:
            link = PortfolioLink.objects.get(user=request.user)
            return Response(PortfolioLinkSerializer(link).data)
        except PortfolioLink.DoesNotExist:
            return Response(None)

    def post(self, request):
        slug = request.data.get('slug', '').strip()
        if not slug or len(slug) < 3:
            return Response({'error': 'Slug must be at least 3 characters'}, status=400)
        if PortfolioLink.objects.filter(slug=slug).exclude(user=request.user).exists():
            return Response({'error': 'Slug already taken'}, status=400)
        # update if exists, create if not
        link, _ = PortfolioLink.objects.update_or_create(
            user=request.user,
            defaults={'slug': slug}
        )
        return Response(PortfolioLinkSerializer(link).data, status=201)

    def delete(self, request):
        try:
            PortfolioLink.objects.get(user=request.user).delete()
            return Response(status=204)
        except PortfolioLink.DoesNotExist:
            return Response({'error': 'No link found'}, status=404)


User = get_user_model()


class PublicPortfolioView(APIView):
    permission_classes = [AllowAny]  # Important: Anyone can view

    def get(self, request, identifier):
        """
        identifier can be either:
        - username
        - custom portfolio slug
        """
        # Step 1: Try finding by custom slug first
        portfolio_link = PortfolioLink.objects.filter(slug=identifier).first()

        if portfolio_link:
            user = portfolio_link.user
        else:
            # Step 2: Fallback to username
            try:
                user = User.objects.get(username=identifier)
            except User.DoesNotExist:
                return Response({"error": "Not found"}, status=404)

        # Return portfolio data
        data = {
            "username": user.username,
            "slug": getattr(user.portfolio_link, 'slug', None),  # custom slug if exists
            "full_name": user.get_full_name() or user.username,
            # Add ALL other fields you want to show on the portfolio page:
            # "bio": user.bio,
            # "profile_picture": user.profile_picture.url if user.profile_picture else None,
            # "projects": [...],
            # "skills": [...],
            # etc.
        }

        return Response(data)


@api_view(['GET'])
@permission_classes([AllowAny])
def get_project_detail(request, pk):
    try:
        project = ProjectsInfo.objects.get(pk=pk)
    except ProjectsInfo.DoesNotExist:
        return Response({'error': 'Project not found'}, status=404)

    try:
        detail = project.detail
        return Response(ProjectDetailSerializer(detail).data)
    except ProjectDetail.DoesNotExist:
        # Return empty shell so frontend can handle gracefully
        return Response({
            'project': pk,
            'intro': '', 'description': '', 'developer': '',
            'github_url': '', 'status': '',
            'tech_stack': [], 'features': [], 'future': [], 'extra_images': [],
        }, status=200)




@api_view(['POST'])
@permission_classes([AllowAny])
def save_project_detail(request, pk):
    try:
        project = ProjectsInfo.objects.get(pk=pk)
    except ProjectsInfo.DoesNotExist:
        return Response({'error': 'Project not found'}, status=404)

    data = request.data.copy()

    # Parse JSON array fields — may arrive as strings or real lists
    for field in ['tech_stack', 'features', 'future', 'extra_images']:
        val = data.get(field, [])
        if isinstance(val, str):
            try:
                parsed = json.loads(val)
                data[field] = parsed if isinstance(parsed, list) else []
            except (json.JSONDecodeError, TypeError):
                data[field] = []
        elif not isinstance(val, list):
            data[field] = []

    # Log incoming data for debugging
    import logging
    logger = logging.getLogger(__name__)
    logger.debug(f"save_project_detail pk={pk} data keys: {list(data.keys())}")
    logger.debug(f"extra_images type={type(data.get('extra_images'))} value={data.get('extra_images')}")

    try:
        detail = project.detail
        serializer = ProjectDetailSerializer(detail, data={**data, 'project': pk}, partial=True)
    except ProjectDetail.DoesNotExist:
        serializer = ProjectDetailSerializer(data={**data, 'project': pk})

    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=200)

    # Return errors clearly so frontend console shows the exact field
    print(f"[save_project_detail] VALIDATION ERRORS: {serializer.errors}")
    return Response({'errors': serializer.errors}, status=400)


@api_view(['GET'])
@permission_classes([AllowAny])
def get_single_project(request, pk):
    try:
        obj = ProjectsInfo.objects.get(pk=pk)
        return Response(ProjectsSerializer(obj).data)
    except ProjectsInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)




# ─────────────────────────────────────────────────────────────────
#  generate_cv (PDF) — fully updated to apply CVPreference
#  Replace your existing generate_cv view with this one.
# ─────────────────────────────────────────────────────────────────

from reportlab.lib.pagesizes import A4, LETTER
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
    KeepTogether, Table, TableStyle, Indenter
)
from reportlab.platypus.paragraph import ParagraphStyle
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
import io

from .models import HomeInfo, AboutInfo, FooterInfo, ProjectsInfo, CVPreference, ALL_SECTIONS


# Map our font_family choices -> ReportLab built-in font names
PDF_FONT_MAP = {
    'Times New Roman': ('Times-Roman', 'Times-Bold', 'Times-Italic'),
    'Helvetica':        ('Helvetica',  'Helvetica-Bold', 'Helvetica-Oblique'),
    'Georgia':           ('Times-Roman', 'Times-Bold', 'Times-Italic'),   # closest built-in fallback
    'Calibri':           ('Helvetica',  'Helvetica-Bold', 'Helvetica-Oblique'),  # closest built-in fallback
    'Garamond':          ('Times-Roman', 'Times-Bold', 'Times-Italic'),   # closest built-in fallback
}


@api_view(['GET'])
@permission_classes([AllowAny])
def generate_cv(request, username):

    home = get_object_or_404(HomeInfo, username=username)

    # ── User selected uploaded CV ─────────────────────────
    if home.cv_mode == "uploaded" and home.cv:
        return redirect(home.cv)

    # ── Route to Word if user selected word format ────────
    if home.cv_format == "word":
        # return generate_word_cv(request, username)
        return generate_word_cv(request._request, username)

    try:
        # ── 0. Load CV preferences (with safe defaults) ──────
        prefs, _ = CVPreference.objects.get_or_create(username=username)

        F, FB, FI = PDF_FONT_MAP.get(prefs.font_family, PDF_FONT_MAP['Times New Roman'])
        BASE_SIZE = prefs.font_size or 12
        LEAD_MULT = prefs.line_spacing or 1.2

        HEAD_COLOR   = colors.HexColor(prefs.heading_color or '#000000')
        TEXT_COLOR   = colors.HexColor(prefs.text_color    or '#000000')
        ACCENT_COLOR = colors.HexColor(prefs.accent_color  or '#000000')

        PAGE = A4 if (prefs.page_size or 'A4') == 'A4' else LETTER

        section_order   = prefs.sections_order or ALL_SECTIONS.copy()
        section_visible = prefs.sections_visible or {s: True for s in ALL_SECTIONS}
        SEC_GAP = prefs.section_padding or 8

        # ── 1. Fetch all form data ────────────────────────────
        home     = HomeInfo.objects.filter(username=username).first()
        about    = AboutInfo.objects.filter(username=username).first()
        footer   = FooterInfo.objects.filter(username=username).first()
        projects = list(ProjectsInfo.objects.filter(username=username).order_by('-id'))

        try:
            from .models import (
                ExperienceInfo, EducationInfo, CertificationInfo,
                AchievementInfo, LanguageSpoken, SoftSkillInfo,
                ProfileSummary, ProfileLinks
            )
            experiences      = list(ExperienceInfo.objects.filter(username=username).order_by('order'))
            educations       = list(EducationInfo.objects.filter(username=username).order_by('order'))
            certifications   = list(CertificationInfo.objects.filter(username=username).order_by('id'))
            achievements     = list(AchievementInfo.objects.filter(username=username).order_by('id'))
            languages_spoken = list(LanguageSpoken.objects.filter(username=username).order_by('id'))
            soft_skills      = list(SoftSkillInfo.objects.filter(username=username).order_by('id'))
            profile_links    = ProfileLinks.objects.filter(username=username).first()
            profile_summary  = ProfileSummary.objects.filter(username=username).first()
        except Exception:
            experiences = educations = certifications = []
            achievements = languages_spoken = soft_skills = []
            profile_links = profile_summary = None

        if not home and not about:
            return Response({'error': 'Portfolio not found'}, status=404)

        # ── 2. Extract all fields safely ──────────────────────
        def g(val, fb=''):
            return str(val).strip() if val else fb

        full_name   = g(home.full_name   if home  else '', 'Your Name')
        skill_title = g(home.skill_title if home  else '', '')
        experience  = g(home.experience  if home  else '', '0')
        s_one       = g(about.s_one      if about else '', '')
        s_two       = g(about.s_two      if about else '', '')
        skill_pack  = g(about.skill_pack if about else '', '')
        edu_legacy  = g(about.education  if about else '', '')

        email        = g(footer.email    if footer else '', '')
        phone        = g(footer.whatsapp if footer else '', '')
        linkedin     = g(footer.linkedin if footer else '', '')
        github       = g(footer.github   if footer else '', '')
        location     = ''
        portfolio_url = f'showcraft.netlify.app/portfolio/{username}'

        if profile_links:
            if g(profile_links.linkedin):      linkedin      = g(profile_links.linkedin)
            if g(profile_links.github):        github        = g(profile_links.github)
            if g(profile_links.phone):         phone         = g(profile_links.phone)
            if g(profile_links.location):      location      = g(profile_links.location)
            if g(profile_links.portfolio_url): portfolio_url = g(profile_links.portfolio_url)

        if phone and not phone.startswith('+'):
            phone = f'+{phone}'

        if profile_summary and g(profile_summary.summary):
            summary_text = g(profile_summary.summary)
        else:
            yrs = experience if experience not in ('0', '') else ''
            summary_text = (
                f"{skill_title}"
                + (f" with {yrs} year{'s' if yrs != '1' else ''} of hands-on experience" if yrs else '')
                + " building full-stack web applications"
                + (f" using {s_one} and {s_two}" if s_one and s_two else '')
                + ". Committed to writing clean, scalable code and contributing to high-impact engineering teams."
            )

        skills_raw = [sk.strip() for sk in skill_pack.split(',') if sk.strip()]

        LANG_SET = {'Python','JavaScript','TypeScript','SQL','HTML','CSS','Java','C++',
                    'C#','PHP','Ruby','Swift','Kotlin','Go','Rust','Dart','R','Bash',
                    'SCSS','Sass','MATLAB'}
        FW_SET   = {'Django REST Framework','React.js','Vue.js','Angular','Next.js',
                    'FastAPI','Flask','Spring Boot','Laravel','Express.js','Tailwind CSS',
                    'Bootstrap','Framer Motion','Redux','Node.js','Vite','React Native',
                    'GraphQL','DRF'}
        DB_SET   = {'PostgreSQL','MySQL','SQLite','MongoDB','Redis','Supabase','Firebase',
                    'Oracle','MSSQL','DynamoDB','Elasticsearch','MariaDB','Cassandra'}
        TOOL_SET = {'Git','GitHub','Postman','VS Code','PyCharm','Docker','Kubernetes',
                    'AWS','Azure','GCP','Nginx','Linux','Jira','Figma','Trello','Heroku',
                    'Railway','Netlify','Vercel','Celery','MinIO','Webpack'}

        langs  = [sk for sk in skills_raw if sk in LANG_SET]
        fws    = [sk for sk in skills_raw if sk in FW_SET]
        dbs    = [sk for sk in skills_raw if sk in DB_SET]
        tools  = [sk for sk in skills_raw if sk in TOOL_SET]
        rest   = [sk for sk in skills_raw if sk not in LANG_SET|FW_SET|DB_SET|TOOL_SET]

        skills_rows = []
        if langs:  skills_rows.append(('Languages',  ', '.join(langs)))
        if fws:    skills_rows.append(('Frameworks', ', '.join(fws)))
        if dbs:    skills_rows.append(('Databases',  ', '.join(dbs)))
        if tools:  skills_rows.append(('Tools',      ', '.join(tools)))
        if rest:   skills_rows.append(('Other',      ', '.join(rest)))

        if soft_skills:
            ss = ', '.join([g(sk.skill) for sk in soft_skills if g(sk.skill)])
            if ss: skills_rows.append(('Soft Skills', ss))

        if languages_spoken:
            ls = ', '.join([
                f'{g(lg.language)} ({g(lg.proficiency)})' if g(lg.proficiency)
                else g(lg.language)
                for lg in languages_spoken if g(lg.language)
            ])
            if ls: skills_rows.append(('Languages Spoken', ls))

        # ── 3. PDF Page Setup (uses prefs for margins + page size) ──
        buffer  = io.BytesIO()
        PAGE_W, PAGE_H = PAGE

        ML = prefs.margin_left   * mm
        MR = prefs.margin_right  * mm
        MT = prefs.margin_top    * mm
        MB = prefs.margin_bottom * mm
        CW = PAGE_W - ML - MR

        doc = SimpleDocTemplate(
            buffer, pagesize=PAGE,
            leftMargin=ML, rightMargin=MR,
            topMargin=MT,  bottomMargin=MB,
            title=f'{full_name} — CV',
            author=full_name,
            subject='Professional Resume',
        )

        BLACK = TEXT_COLOR
        DARK  = TEXT_COLOR

        BS = BASE_SIZE
        LS = BS * LEAD_MULT

        # ── 6. Styles (sizes/colors driven by prefs) ──────────
        def mk(name, **kw):
            base = dict(
                fontName=F,
                fontSize=BS - 1.85,
                textColor=BLACK,
                leading=LS,
                spaceAfter=0,
                spaceBefore=0,
                alignment=TA_JUSTIFY,
                splitLongWords=True,
                allowWidows=1,
                allowOrphans=1,
                justifyBreaks=1,
            )
            base.update(kw)
            return ParagraphStyle(name, **base)

        sName = mk('nm', fontName=FB, fontSize=BS + 8,
                   alignment=TA_CENTER, leading=BS + 12, textColor=HEAD_COLOR)

        sContact = mk('ct', fontSize=BS - 1, alignment=TA_CENTER,
                      leading=BS + 3, textColor=DARK)

        sSec = mk('sc', fontName=FB, fontSize=BS - 0.7, leading=BS + 2,
                   alignment=TA_LEFT, textColor=HEAD_COLOR)

        sBody = mk('bd', fontName=F, fontSize=BS - 1.85, leading=LS,
                    alignment=TA_JUSTIFY, textColor=BLACK)

        sBold = mk('bl', fontName=FB, fontSize=BS - 1.6, leading=LS - 0.2,
                    alignment=TA_LEFT, textColor=BLACK)

        sPlain = mk('pl', fontName=F, fontSize=BS - 2, leading=LS - 0.7,
                     alignment=TA_LEFT, textColor=DARK)

        sUrl = mk('url', fontName=F, fontSize=BS - 2.3, leading=LS - 1.4,
                   alignment=TA_LEFT, textColor=BLACK)

        sSkillLabel = mk('skl', fontName=FB, fontSize=BS - 1.9, leading=LS - 0.4,
                          alignment=TA_LEFT, textColor=BLACK)

        sSkillVal = mk('skv', fontName=F, fontSize=BS - 1.9, leading=LS - 0.4,
                        alignment=TA_LEFT, textColor=BLACK)

        sBullet = mk('bu', fontName=F, fontSize=BS - 1.9, leading=LS + 0.1,
                      leftIndent=18, firstLineIndent=-10, spaceAfter=2,
                      alignment=TA_LEFT, textColor=BLACK)

        # ── 7. Helpers ────────────────────────────────────────
        def sp(pts):
            return Spacer(1, pts)

        def rule():
            return HRFlowable(
                width='100%', thickness=0.75, color=ACCENT_COLOR,
                lineCap='square', spaceBefore=1, spaceAfter=6,
            )

        def section(title):
            return [Spacer(1, SEC_GAP), Paragraph(title.upper(), sSec), rule()]

        def bul(text):
            text = text.strip().lstrip('-').strip()
            return Paragraph(f'• {text}', sBullet)

        # ── 8. Build Content ─────────────────────────────────
        story = []

        # HEADER
        story.append(Paragraph(full_name.upper(), sName))
        story.append(Spacer(1, 3))

        contact_parts = []
        if location: contact_parts.append(location)
        if phone:    contact_parts.append(phone)
        if email:
            contact_parts.append(f'<link href="mailto:{email}" color="black">{email}</link>')
        if linkedin:
            url = linkedin.strip()
            if not url.startswith('http'): url = f'https://{url}'
            contact_parts.append(f'<link href="{url}" color="black">LinkedIn</link>')
        if github:
            url = github.strip()
            if not url.startswith('http'): url = f'https://{url}'
            contact_parts.append(f'<link href="{url}" color="black">GitHub</link>')
        if portfolio_url:
            url = portfolio_url.strip()
            if not url.startswith('http'): url = f'https://{url}'
            contact_parts.append(f'<link href="{url}" color="black">Portfolio</link>')

        story.append(Paragraph(' | '.join(contact_parts), sContact))
        story.append(Spacer(1, 10))

        # ── Section renderers (called dynamically in order) ────
        def render_summary():
            out = section('Summary')
            out.append(Paragraph(summary_text, sBody))
            return out

        def render_education():
            out = []
            if educations:
                out += section('EDUCATION')
                for edu in educations:
                    block = [Paragraph(f'<b>{g(edu.degree)}</b>', sBold)]
                    meta = []
                    if g(edu.university): meta.append(g(edu.university))
                    if g(edu.location):   meta.append(g(edu.location))
                    if g(edu.session):    meta.append(g(edu.session))
                    if meta: block.append(Paragraph(' | '.join(meta), sPlain))
                    block.append(Spacer(1, 5))
                    out.append(KeepTogether(block))
            elif edu_legacy:
                out += section('EDUCATION')
                out.append(Paragraph(edu_legacy, sBody))
                out.append(Spacer(1, 5))
            return out

        def render_experience():
            out = []
            if experiences:
                out += section('EXPERIENCE')
                for exp in experiences:
                    block = [Paragraph(f'<b>{g(exp.title)}</b> - {g(exp.company)}', sBold)]
                    meta = []
                    if g(exp.dates):    meta.append(g(exp.dates))
                    if g(exp.location): meta.append(g(exp.location))
                    if meta: block.append(Paragraph(' | '.join(meta), sPlain))
                    block.append(Spacer(1, 4))
                    if g(exp.bullets):
                        for line in g(exp.bullets).split('\n'):
                            if line.strip(): block.append(bul(line))
                    block.append(Spacer(1, 5))
                    out.append(KeepTogether(block))
            return out

        def render_skills():
            out = []
            if skills_rows:
                out += section('SKILLS')
                indent_amount = 8
                out.append(Indenter(left=indent_amount))
                for label, value in skills_rows:
                    row = [[Paragraph(f'<b>{label}:</b>', sSkillLabel), Paragraph(value, sSkillVal)]]
                    table = Table(row, colWidths=[(CW - indent_amount) * 0.20, (CW - indent_amount) * 0.80])
                    table.setStyle(TableStyle([
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), 0),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                        ('TOPPADDING', (0, 0), (-1, -1), 0.5),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 1.5),
                    ]))
                    out.append(table)
                out.append(Indenter(left=-indent_amount))
                out.append(Spacer(1, 4))
            return out

        def render_certifications():
            out = []
            if certifications:
                out += section('CERTIFICATIONS')
                for cert in certifications:
                    block = []
                    cert_line = g(cert.name)
                    if g(cert.date): cert_line += f' | {g(cert.date)}'
                    block.append(Paragraph(cert_line, sBold))
                    if g(cert.issuer): block.append(Paragraph(g(cert.issuer), sPlain))
                    if g(cert.verify_url): block.append(Paragraph(f'Verify: {g(cert.verify_url)}', sPlain))
                    block.append(Spacer(1, 5))
                    out.append(KeepTogether(block))
            return out

        def render_achievements():
            out = []
            if achievements:
                out += section('HONORS & ACHIEVEMENTS')
                for ach in achievements:
                    block = [Paragraph(f'<b>{g(ach.title)}</b>', sBold), Spacer(1, 3)]
                    if g(ach.description):
                        for line in g(ach.description).split('\n'):
                            if line.strip(): block.append(bul(line))
                    block.append(Spacer(1, 5))
                    out.append(KeepTogether(block))
            return out

        def render_projects():
            out = []
            if projects:
                out += section('PROJECTS')
                for p in projects:
                    block = [Paragraph(f'<b>{g(p.p_name)}</b>', sBold)]
                    p_url = g(p.p_url)
                    if p_url:
                        clean_url = p_url.strip()
                        if not clean_url.startswith('http'): clean_url = f'https://{clean_url}'
                        label = 'GitHub' if 'github' in clean_url.lower() else 'Live'
                        block.append(Paragraph(
                            f'<link href="{clean_url}" color="black">{label}: {clean_url}</link>', sUrl
                        ))
                    block.append(Spacer(1, 4))
                    desc_text = ''
                    try:
                        from .models import ProjectDetail
                        detail = ProjectDetail.objects.filter(project=p).first()
                        if detail and g(detail.description):
                            desc_text = g(detail.description)
                    except Exception:
                        pass
                    if desc_text:
                        for line in desc_text.split('\n'):
                            if line.strip(): block.append(bul(line))
                    block.append(Spacer(1, 7))
                    out.append(KeepTogether(block))
            return out

        RENDERERS = {
            'summary':        render_summary,
            'education':      render_education,
            'experience':     render_experience,
            'skills':         render_skills,
            'certifications': render_certifications,
            'achievements':   render_achievements,
            'projects':       render_projects,
        }

        # ── Render sections in user-defined order, skip hidden ──
        for key in section_order:
            if section_visible.get(key, True) and key in RENDERERS:
                story += RENDERERS[key]()

        # ── 9. Build PDF ──────────────────────────────────────
        doc.build(story)
        buffer.seek(0)

        # Custom file name (falls back to full name)
        custom_name = (prefs.file_name or '').strip()
        if custom_name:
            safe_name = custom_name.replace(' ', '_').replace('/', '_')
        else:
            safe_name = full_name.replace(' ', '_').replace('/', '_')
        filename = f'{safe_name}.pdf' if not safe_name.lower().endswith('.pdf') else safe_name

        response = HttpResponse(buffer.read(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        response['Cache-Control']       = 'no-cache'
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({'error': f'CV generation failed: {str(e)}'}, status=500)




# ─────────────────────────────────────────────────────────────────
#  generate_word_cv — fully updated to apply CVPreference
#  Replace your existing generate_word_cv view with this one.
# ─────────────────────────────────────────────────────────────────



@api_view(['GET'])
@permission_classes([AllowAny])
def generate_word_cv(request, username):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:

        import logging
        logger = logging.getLogger(__name__)

        prefs, _ = CVPreference.objects.get_or_create(username=username)
        logger.error(f"WORD CV DEBUG: prefs loaded for {username}")

        # ── 0. Load preferences ────────────────────────────
        # prefs, _ = CVPreference.objects.get_or_create(username=username)

        FONT_FAMILY = prefs.font_family or 'Times New Roman'
        BASE_SIZE   = prefs.font_size or 12
        LINE_MULT   = prefs.line_spacing or 1.2

        #HEAD_RGB = RGBColor.from_string((prefs.heading_color or '#000000').lstrip('#'))
        #TEXT_RGB = RGBColor.from_string((prefs.text_color    or '#000000').lstrip('#'))
       #ACCENT_HEX = (prefs.accent_color or '#000000').lstrip('#')

       # CORRECT
        def hex_to_rgb(hex_color):
           h = (hex_color or '#000000').lstrip('#')
           return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        HEAD_RGB   = hex_to_rgb(prefs.heading_color)

        TEXT_RGB   = hex_to_rgb(prefs.text_color)
        ACCENT_HEX = (prefs.accent_color or '#000000').lstrip('#')

        section_order   = prefs.sections_order or ALL_SECTIONS.copy()
        section_visible = prefs.sections_visible or {s: True for s in ALL_SECTIONS}
        SEC_GAP = Pt(prefs.section_padding or 8)

        # ── 1. Fetch data (identical to PDF) ──────────────────
        home     = HomeInfo.objects.filter(username=username).first()
        about    = AboutInfo.objects.filter(username=username).first()
        footer   = FooterInfo.objects.filter(username=username).first()
        projects = list(ProjectsInfo.objects.filter(username=username).order_by('-id'))

        if not home:
            return Response({'error': 'Portfolio not found'}, status=404)

        try:
            from .models import (
                ExperienceInfo, EducationInfo, CertificationInfo,
                AchievementInfo, LanguageSpoken, SoftSkillInfo,
                ProfileSummary, ProfileLinks
            )
            experiences      = list(ExperienceInfo.objects.filter(username=username).order_by('order'))
            educations       = list(EducationInfo.objects.filter(username=username).order_by('order'))
            certifications   = list(CertificationInfo.objects.filter(username=username).order_by('id'))
            achievements     = list(AchievementInfo.objects.filter(username=username).order_by('id'))
            languages_spoken = list(LanguageSpoken.objects.filter(username=username).order_by('id'))
            soft_skills      = list(SoftSkillInfo.objects.filter(username=username).order_by('id'))
            profile_links    = ProfileLinks.objects.filter(username=username).first()
            profile_summary  = ProfileSummary.objects.filter(username=username).first()
        except Exception:
            experiences = educations = certifications = []
            achievements = languages_spoken = soft_skills = []
            profile_links = profile_summary = None

        def g(val, fb=''):
            return str(val).strip() if val else fb

        full_name   = g(home.full_name,   'Your Name')
        skill_title = g(home.skill_title, '')
        experience  = g(home.experience,  '0')
        s_one       = g(about.s_one      if about else '', '')
        s_two       = g(about.s_two      if about else '', '')
        skill_pack  = g(about.skill_pack if about else '', '')
        edu_legacy  = g(about.education  if about else '', '')

        email         = g(footer.email    if footer else '', '')
        phone         = g(footer.whatsapp if footer else '', '')
        linkedin      = g(footer.linkedin if footer else '', '')
        github        = g(footer.github   if footer else '', '')
        location      = ''
        portfolio_url = f'showcraft.netlify.app/portfolio/{username}'

        if profile_links:
            if g(profile_links.linkedin):      linkedin      = g(profile_links.linkedin)
            if g(profile_links.github):        github        = g(profile_links.github)
            if g(profile_links.phone):         phone         = g(profile_links.phone)
            if g(profile_links.location):      location      = g(profile_links.location)
            if g(profile_links.portfolio_url): portfolio_url = g(profile_links.portfolio_url)

        if phone and not phone.startswith('+'):
            phone = f'+{phone}'

        if profile_summary and g(profile_summary.summary):
            summary_text = g(profile_summary.summary)
        else:
            yrs = experience if experience not in ('0', '') else ''
            summary_text = (
                f"{skill_title}"
                + (f" with {yrs} year{'s' if yrs != '1' else ''} of hands-on experience" if yrs else '')
                + " building full-stack web applications"
                + (f" using {s_one} and {s_two}" if s_one and s_two else '')
                + ". Committed to writing clean, scalable code and contributing to high-impact engineering teams."
            )

        skills_raw = [sk.strip() for sk in skill_pack.split(',') if sk.strip()]
        LANG_SET = {'Python','JavaScript','TypeScript','SQL','HTML','CSS','Java','C++',
                    'C#','PHP','Ruby','Swift','Kotlin','Go','Rust','Dart','R','Bash',
                    'SCSS','Sass','MATLAB'}
        FW_SET   = {'Django REST Framework','React.js','Vue.js','Angular','Next.js',
                    'FastAPI','Flask','Spring Boot','Laravel','Express.js','Tailwind CSS',
                    'Bootstrap','Framer Motion','Redux','Node.js','Vite','React Native',
                    'GraphQL','DRF'}
        DB_SET   = {'PostgreSQL','MySQL','SQLite','MongoDB','Redis','Supabase','Firebase',
                    'Oracle','MSSQL','DynamoDB','Elasticsearch','MariaDB','Cassandra'}
        TOOL_SET = {'Git','GitHub','Postman','VS Code','PyCharm','Docker','Kubernetes',
                    'AWS','Azure','GCP','Nginx','Linux','Jira','Figma','Trello','Heroku',
                    'Railway','Netlify','Vercel','Celery','MinIO','Webpack'}

        langs  = [sk for sk in skills_raw if sk in LANG_SET]
        fws    = [sk for sk in skills_raw if sk in FW_SET]
        dbs    = [sk for sk in skills_raw if sk in DB_SET]
        tools  = [sk for sk in skills_raw if sk in TOOL_SET]
        rest   = [sk for sk in skills_raw if sk not in LANG_SET|FW_SET|DB_SET|TOOL_SET]

        skills_rows = []
        if langs:  skills_rows.append(('Languages',  ', '.join(langs)))
        if fws:    skills_rows.append(('Frameworks', ', '.join(fws)))
        if dbs:    skills_rows.append(('Databases',  ', '.join(dbs)))
        if tools:  skills_rows.append(('Tools',      ', '.join(tools)))
        if rest:   skills_rows.append(('Other',      ', '.join(rest)))
        if soft_skills:
            ss = ', '.join([g(sk.skill) for sk in soft_skills if g(sk.skill)])
            if ss: skills_rows.append(('Soft Skills', ss))
        if languages_spoken:
            ls = ', '.join([
                f'{g(lg.language)} ({g(lg.proficiency)})' if g(lg.proficiency) else g(lg.language)
                for lg in languages_spoken if g(lg.language)
            ])
            if ls: skills_rows.append(('Languages Spoken', ls))

        # ── 3. Build the Word document ────────────────────────
        doc = Document()

        for sec in doc.sections:
            sec.top_margin    = Cm(prefs.margin_top    / 10)   # mm -> cm
            sec.bottom_margin = Cm(prefs.margin_bottom / 10)
            sec.left_margin   = Cm(prefs.margin_left   / 10)
            sec.right_margin  = Cm(prefs.margin_right  / 10)

            # Page size
            from docx.shared import Mm
            if (prefs.page_size or 'A4') == 'Letter':
                sec.page_width  = Mm(215.9)
                sec.page_height = Mm(279.4)
            else:
                sec.page_width  = Mm(210)
                sec.page_height = Mm(297)

        normal = doc.styles['Normal']
        normal.font.name = FONT_FAMILY
        normal.font.size = Pt(BASE_SIZE)
        normal.font.color.rgb = TEXT_RGB
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after  = Pt(0)
        normal.paragraph_format.line_spacing = LINE_MULT

        def add_thick_rule(paragraph):
            pPr = paragraph._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '12')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), ACCENT_HEX)
            pBdr.append(bottom)
            pPr.append(pBdr)

        def add_section_heading(title):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = SEC_GAP
            p.paragraph_format.space_after  = Pt(0)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.name = FONT_FAMILY
            run.font.size = Pt(BASE_SIZE - 0.7)
            run.font.color.rgb = HEAD_RGB
            add_thick_rule(p)

        def add_body(text, bold=False):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.line_spacing = LINE_MULT
            run = p.add_run(text)
            run.bold = bold
            run.font.name = FONT_FAMILY
            run.font.size = Pt(BASE_SIZE)
            run.font.color.rgb = TEXT_RGB
            return p

        def add_bullet(text):
            text = text.strip().lstrip('-').strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.line_spacing = LINE_MULT
            p.paragraph_format.left_indent       = Cm(0.45)
            p.paragraph_format.first_line_indent = Cm(-0.25)
            run = p.add_run(f'- {text}')
            run.font.name = FONT_FAMILY
            run.font.size = Pt(BASE_SIZE)
            run.font.color.rgb = TEXT_RGB

        def add_meta(text):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(text)
            run.font.name = FONT_FAMILY
            run.font.size = Pt(max(BASE_SIZE - 2, 8))
            run.font.color.rgb = TEXT_RGB
            return p

        def small_gap():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(4)

        # ── 4. Header ──────────────────────────────────────────
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.paragraph_format.space_after = Pt(2)
        name_run = name_p.add_run(full_name.upper())
        name_run.bold = True
        name_run.font.name = FONT_FAMILY
        name_run.font.size = Pt(BASE_SIZE + 10)
        name_run.font.color.rgb = HEAD_RGB

        contact_parts = []
        if location:      contact_parts.append(location)
        if phone:         contact_parts.append(phone)
        if email:         contact_parts.append(email)
        if linkedin:      contact_parts.append(linkedin)
        if github:        contact_parts.append(github)
        if portfolio_url: contact_parts.append(portfolio_url)

        if contact_parts:
            cp = doc.add_paragraph(' | '.join(contact_parts))
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(8)
            for run in cp.runs:
                run.font.name = FONT_FAMILY
                run.font.size = Pt(max(BASE_SIZE - 2, 8))
                run.font.color.rgb = TEXT_RGB

        rule_p = doc.add_paragraph()
        rule_p.paragraph_format.space_before = Pt(2)
        rule_p.paragraph_format.space_after  = Pt(4)
        add_thick_rule(rule_p)

        # ── Section renderers ────────────────────────────────
        def render_summary():
            add_section_heading('Summary')
            add_body(summary_text)

        def render_education():
            if educations:
                add_section_heading('Education')
                for edu in educations:
                    add_body(g(edu.degree), bold=True)
                    meta = []
                    if g(edu.university): meta.append(g(edu.university))
                    if g(edu.location):   meta.append(g(edu.location))
                    if g(edu.session):    meta.append(g(edu.session))
                    if meta: add_meta(' | '.join(meta))
                    small_gap()
            elif edu_legacy:
                add_section_heading('Education')
                add_body(edu_legacy)
                small_gap()

        def render_experience():
            if experiences:
                add_section_heading('Experience')
                for exp in experiences:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(1)
                    r1 = p.add_run(g(exp.title)); r1.bold = True
                    r1.font.name = FONT_FAMILY; r1.font.size = Pt(BASE_SIZE); r1.font.color.rgb = TEXT_RGB
                    r2 = p.add_run(f' - {g(exp.company)}')
                    r2.font.name = FONT_FAMILY; r2.font.size = Pt(BASE_SIZE); r2.font.color.rgb = TEXT_RGB

                    meta = []
                    if g(exp.dates):    meta.append(g(exp.dates))
                    if g(exp.location): meta.append(g(exp.location))
                    if meta: add_meta(' | '.join(meta))

                    if g(exp.bullets):
                        for line in g(exp.bullets).split('\n'):
                            if line.strip(): add_bullet(line)
                    small_gap()

        def render_skills():
            if skills_rows:
                add_section_heading('Skills')
                for label, value in skills_rows:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after  = Pt(1.5)
                    p.paragraph_format.left_indent  = Cm(0.45)
                    r1 = p.add_run(f'{label}: '); r1.bold = True
                    r1.font.name = FONT_FAMILY; r1.font.size = Pt(BASE_SIZE); r1.font.color.rgb = TEXT_RGB
                    r2 = p.add_run(value)
                    r2.font.name = FONT_FAMILY; r2.font.size = Pt(BASE_SIZE); r2.font.color.rgb = TEXT_RGB

        def render_certifications():
            if certifications:
                add_section_heading('Certifications')
                for cert in certifications:
                    cert_line = g(cert.name)
                    if g(cert.date): cert_line += f' | {g(cert.date)}'
                    add_body(cert_line, bold=True)
                    if g(cert.issuer):     add_meta(g(cert.issuer))
                    if g(cert.verify_url): add_meta(f'Verify: {g(cert.verify_url)}')
                    small_gap()

        def render_achievements():
            if achievements:
                add_section_heading('Honours & Achievements')
                for ach in achievements:
                    add_body(g(ach.title), bold=True)
                    if g(ach.description):
                        for line in g(ach.description).split('\n'):
                            if line.strip(): add_bullet(line)
                    small_gap()

        def render_projects():
            if projects:
                add_section_heading('Projects')
                for proj in projects:
                    add_body(g(proj.p_name), bold=True)
                    if g(proj.p_url):
                        url = g(proj.p_url).strip()
                        label = 'GitHub' if 'github' in url.lower() else 'Live'
                        add_meta(f'{label}: {url}')
                    try:
                        from .models import ProjectDetail
                        detail = ProjectDetail.objects.filter(project=proj).first()
                        if detail and g(detail.description):
                            for line in g(detail.description).split('\n'):
                                if line.strip(): add_bullet(line)
                    except Exception:
                        pass
                    small_gap()

        RENDERERS = {
            'summary':        render_summary,
            'education':      render_education,
            'experience':     render_experience,
            'skills':         render_skills,
            'certifications': render_certifications,
            'achievements':   render_achievements,
            'projects':       render_projects,
        }

        for key in section_order:
            if section_visible.get(key, True) and key in RENDERERS:
                RENDERERS[key]()

                # ── 13. Stream response ───────────────────────────────
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                safe_name = full_name.replace(' ', '_').replace('/', '_')
                filename = f'{safe_name}_CV.docx'

                # Use buffer.getvalue() to pass the raw bytes
                response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                )

            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            response['Cache-Control'] = 'no-cache'
        return response

    except Exception as e:
         import traceback
         traceback.print_exc()
         return Response({'error': f'Word CV generation failed: {str(e)}'}, status=500)



# ─────────────────────────────────────────────────────────────────
#  WORD CV GENERATOR  (same data, same sections, .docx output)
# ─────────────────────────────────────────────────────────────────



#
# @api_view(['GET'])
# @permission_classes([AllowAny])
# def generate_word_cv(request, username):
#     """
#     Generates a .docx CV using the exact same data as generate_cv().
#     Matches the same sections: Summary, Education, Experience,
#     Skills, Certifications, Achievements, Projects.
#     Typography: Times New Roman 12pt, bold headings, thick rule.
#     """
#
#
#     try:
#         # ── 1. Fetch all data (identical to generate_cv) ─────
#         home = HomeInfo.objects.filter(username=username).first()
#         about = AboutInfo.objects.filter(username=username).first()
#         footer = FooterInfo.objects.filter(username=username).first()
#         projects = list(ProjectsInfo.objects.filter(username=username).order_by('-id'))
#
#         if not home:
#             return Response({'error': 'Portfolio not found'}, status=404)
#
#         try:
#             from .models import (
#                 ExperienceInfo, EducationInfo, CertificationInfo,
#                 AchievementInfo, LanguageSpoken, SoftSkillInfo,
#                 ProfileSummary, ProfileLinks
#             )
#             experiences = list(ExperienceInfo.objects.filter(username=username).order_by('order'))
#             educations = list(EducationInfo.objects.filter(username=username).order_by('order'))
#             certifications = list(CertificationInfo.objects.filter(username=username).order_by('id'))
#             achievements = list(AchievementInfo.objects.filter(username=username).order_by('id'))
#             languages_spoken = list(LanguageSpoken.objects.filter(username=username).order_by('id'))
#             soft_skills = list(SoftSkillInfo.objects.filter(username=username).order_by('id'))
#             profile_links = ProfileLinks.objects.filter(username=username).first()
#             profile_summary = ProfileSummary.objects.filter(username=username).first()
#         except Exception:
#             experiences = educations = certifications = []
#             achievements = languages_spoken = soft_skills = []
#             profile_links = profile_summary = None
#
#         # ── 2. Extract fields (identical to generate_cv) ─────
#         def g(val, fb=''):
#             return str(val).strip() if val else fb
#
#         full_name = g(home.full_name, 'Your Name')
#         skill_title = g(home.skill_title, '')
#         experience = g(home.experience, '0')
#         s_one = g(about.s_one if about else '', '')
#         s_two = g(about.s_two if about else '', '')
#         skill_pack = g(about.skill_pack if about else '', '')
#         edu_legacy = g(about.education if about else '', '')
#
#         email = g(footer.email if footer else '', '')
#         phone = g(footer.whatsapp if footer else '', '')
#         linkedin = g(footer.linkedin if footer else '', '')
#         github = g(footer.github if footer else '', '')
#         location = ''
#         portfolio_url = f'showcraft.netlify.app/portfolio/{username}'
#
#         if profile_links:
#             if g(profile_links.linkedin):      linkedin = g(profile_links.linkedin)
#             if g(profile_links.github):        github = g(profile_links.github)
#             if g(profile_links.phone):         phone = g(profile_links.phone)
#             if g(profile_links.location):      location = g(profile_links.location)
#             if g(profile_links.portfolio_url): portfolio_url = g(profile_links.portfolio_url)
#
#         if phone and not phone.startswith('+'):
#             phone = f'+{phone}'
#
#         if profile_summary and g(profile_summary.summary):
#             summary_text = g(profile_summary.summary)
#         else:
#             yrs = experience if experience not in ('0', '') else ''
#             summary_text = (
#                     f"{skill_title}"
#                     + (f" with {yrs} year{'s' if yrs != '1' else ''} of hands-on experience" if yrs else '')
#                     + " building full-stack web applications"
#                     + (f" using {s_one} and {s_two}" if s_one and s_two else '')
#                     + ". Committed to writing clean, scalable code and contributing to high-impact engineering teams."
#             )
#
#         skills_raw = [sk.strip() for sk in skill_pack.split(',') if sk.strip()]
#         LANG_SET = {'Python', 'JavaScript', 'TypeScript', 'SQL', 'HTML', 'CSS', 'Java', 'C++',
#                     'C#', 'PHP', 'Ruby', 'Swift', 'Kotlin', 'Go', 'Rust', 'Dart', 'R', 'Bash',
#                     'SCSS', 'Sass', 'MATLAB'}
#         FW_SET = {'Django REST Framework', 'React.js', 'Vue.js', 'Angular', 'Next.js',
#                   'FastAPI', 'Flask', 'Spring Boot', 'Laravel', 'Express.js', 'Tailwind CSS',
#                   'Bootstrap', 'Framer Motion', 'Redux', 'Node.js', 'Vite', 'React Native',
#                   'GraphQL', 'DRF'}
#         DB_SET = {'PostgreSQL', 'MySQL', 'SQLite', 'MongoDB', 'Redis', 'Supabase', 'Firebase',
#                   'Oracle', 'MSSQL', 'DynamoDB', 'Elasticsearch', 'MariaDB', 'Cassandra'}
#         TOOL_SET = {'Git', 'GitHub', 'Postman', 'VS Code', 'PyCharm', 'Docker', 'Kubernetes',
#                     'AWS', 'Azure', 'GCP', 'Nginx', 'Linux', 'Jira', 'Figma', 'Trello', 'Heroku',
#                     'Railway', 'Netlify', 'Vercel', 'Celery', 'MinIO', 'Webpack'}
#
#         langs = [sk for sk in skills_raw if sk in LANG_SET]
#         fws = [sk for sk in skills_raw if sk in FW_SET]
#         dbs = [sk for sk in skills_raw if sk in DB_SET]
#         tools = [sk for sk in skills_raw if sk in TOOL_SET]
#         rest = [sk for sk in skills_raw if sk not in LANG_SET | FW_SET | DB_SET | TOOL_SET]
#
#         skills_rows = []
#         if langs:  skills_rows.append(('Languages', ', '.join(langs)))
#         if fws:    skills_rows.append(('Frameworks', ', '.join(fws)))
#         if dbs:    skills_rows.append(('Databases', ', '.join(dbs)))
#         if tools:  skills_rows.append(('Tools', ', '.join(tools)))
#         if rest:   skills_rows.append(('Other', ', '.join(rest)))
#         if soft_skills:
#             ss = ', '.join([g(sk.skill) for sk in soft_skills if g(sk.skill)])
#             if ss: skills_rows.append(('Soft Skills', ss))
#         if languages_spoken:
#             ls = ', '.join([
#                 f'{g(lg.language)} ({g(lg.proficiency)})' if g(lg.proficiency) else g(lg.language)
#                 for lg in languages_spoken if g(lg.language)
#             ])
#             if ls: skills_rows.append(('Languages Spoken', ls))
#
#         # ── 3. Build the Word document ────────────────────────
#         doc = Document()
#
#         # Page margins — match PDF (13mm each side)
#         for sec in doc.sections:
#             sec.top_margin = Cm(1.2)
#             sec.bottom_margin = Cm(1.2)
#             sec.left_margin = Cm(1.3)
#             sec.right_margin = Cm(1.3)
#
#         # Remove default paragraph spacing globally
#         from docx.styles.style import _ParagraphStyle
#         normal = doc.styles['Normal']
#         normal.font.name = 'Times New Roman'
#         normal.font.size = Pt(12)
#         normal.paragraph_format.space_before = Pt(0)
#         normal.paragraph_format.space_after = Pt(0)
#
#         # ── Helpers ───────────────────────────────────────────
#
#         def add_thick_rule(paragraph):
#             """Add a thick bottom border under a paragraph (like PDF rule())."""
#             pPr = paragraph._p.get_or_add_pPr()
#             pBdr = OxmlElement('w:pBdr')
#             bottom = OxmlElement('w:bottom')
#             bottom.set(qn('w:val'), 'single')
#             bottom.set(qn('w:sz'), '12')  # 1.5pt thick
#             bottom.set(qn('w:space'), '1')
#             bottom.set(qn('w:color'), '000000')
#             pBdr.append(bottom)
#             pPr.append(pBdr)
#
#         def add_section_heading(title):
#             """Bold uppercase heading with thick rule beneath — matches PDF section()."""
#             p = doc.add_paragraph()
#             p.paragraph_format.space_before = Pt(8)
#             p.paragraph_format.space_after = Pt(0)
#             run = p.add_run(title.upper())
#             run.bold = True
#             run.font.name = 'Times New Roman'
#             run.font.size = Pt(11.3)
#             add_thick_rule(p)
#
#         def add_body(text, bold=False, indent=False):
#             """Standard 12pt body paragraph."""
#             p = doc.add_paragraph()
#             p.paragraph_format.space_before = Pt(0)
#             p.paragraph_format.space_after = Pt(1)
#             if indent:
#                 p.paragraph_format.left_indent = Cm(0.45)
#             run = p.add_run(text)
#             run.bold = bold
#             run.font.name = 'Times New Roman'
#             run.font.size = Pt(12)
#             return p
#
#         def add_bullet(text):
#             """Dash bullet — matches PDF bul()."""
#             text = text.strip().lstrip('-').strip()
#             p = doc.add_paragraph()
#             p.paragraph_format.space_before = Pt(0)
#             p.paragraph_format.space_after = Pt(2)
#             p.paragraph_format.left_indent = Cm(0.45)
#             p.paragraph_format.first_line_indent = Cm(-0.25)
#             run = p.add_run(f'- {text}')
#             run.font.name = 'Times New Roman'
#             run.font.size = Pt(12)
#
#         def add_meta(text):
#             """Smaller plain line for dates/location/issuer."""
#             p = doc.add_paragraph()
#             p.paragraph_format.space_before = Pt(0)
#             p.paragraph_format.space_after = Pt(1)
#             run = p.add_run(text)
#             run.font.name = 'Times New Roman'
#             run.font.size = Pt(10)
#             return p
#
#         def small_gap():
#             p = doc.add_paragraph()
#             p.paragraph_format.space_before = Pt(0)
#             p.paragraph_format.space_after = Pt(4)
#
#         # ── 4. Header: Name ───────────────────────────────────
#         name_p = doc.add_paragraph()
#         name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
#         name_p.paragraph_format.space_after = Pt(2)
#         name_run = name_p.add_run(full_name.upper())
#         name_run.bold = True
#         name_run.font.name = 'Times New Roman'
#         name_run.font.size = Pt(22)
#
#         # ── 5. Contact line ───────────────────────────────────
#         contact_parts = []
#         if location:      contact_parts.append(location)
#         if phone:         contact_parts.append(phone)
#         if email:         contact_parts.append(email)
#         if linkedin:      contact_parts.append(linkedin)
#         if github:        contact_parts.append(github)
#         if portfolio_url: contact_parts.append(portfolio_url)
#
#         if contact_parts:
#             cp = doc.add_paragraph(' | '.join(contact_parts))
#             cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
#             cp.paragraph_format.space_after = Pt(8)
#             for run in cp.runs:
#                 run.font.name = 'Times New Roman'
#                 run.font.size = Pt(10)
#
#         # Thick rule after header (matches PDF)
#         rule_p = doc.add_paragraph()
#         rule_p.paragraph_format.space_before = Pt(2)
#         rule_p.paragraph_format.space_after = Pt(4)
#         add_thick_rule(rule_p)
#
#         # ── 6. Summary ────────────────────────────────────────
#         add_section_heading('Summary')
#         add_body(summary_text)
#
#         # ── 7. Education ──────────────────────────────────────
#         if educations:
#             add_section_heading('Education')
#             for edu in educations:
#                 add_body(g(edu.degree), bold=True)
#                 meta = []
#                 if g(edu.university): meta.append(g(edu.university))
#                 if g(edu.location):   meta.append(g(edu.location))
#                 if g(edu.session):    meta.append(g(edu.session))
#                 if meta: add_meta(' | '.join(meta))
#                 small_gap()
#         elif edu_legacy:
#             add_section_heading('Education')
#             add_body(edu_legacy)
#             small_gap()
#
#         # ── 8. Experience ─────────────────────────────────────
#         if experiences:
#             add_section_heading('Experience')
#             for exp in experiences:
#                 # Title + company
#                 p = doc.add_paragraph()
#                 p.paragraph_format.space_before = Pt(0)
#                 p.paragraph_format.space_after = Pt(1)
#                 r1 = p.add_run(g(exp.title))
#                 r1.bold = True
#                 r1.font.name = 'Times New Roman'
#                 r1.font.size = Pt(12)
#                 r2 = p.add_run(f' - {g(exp.company)}')
#                 r2.font.name = 'Times New Roman'
#                 r2.font.size = Pt(12)
#
#                 # Dates + location
#                 meta = []
#                 if g(exp.dates):    meta.append(g(exp.dates))
#                 if g(exp.location): meta.append(g(exp.location))
#                 if meta: add_meta(' | '.join(meta))
#
#                 # Bullets
#                 if g(exp.bullets):
#                     for line in g(exp.bullets).split('\n'):
#                         if line.strip():
#                             add_bullet(line)
#                 small_gap()
#
#         # ── 9. Skills ─────────────────────────────────────────
#         if skills_rows:
#             add_section_heading('Skills')
#             for label, value in skills_rows:
#                 p = doc.add_paragraph()
#                 p.paragraph_format.space_before = Pt(0)
#                 p.paragraph_format.space_after = Pt(1.5)
#                 p.paragraph_format.left_indent = Cm(0.45)
#                 r1 = p.add_run(f'{label}: ')
#                 r1.bold = True
#                 r1.font.name = 'Times New Roman'
#                 r1.font.size = Pt(12)
#                 r2 = p.add_run(value)
#                 r2.font.name = 'Times New Roman'
#                 r2.font.size = Pt(12)
#
#         # ── 10. Certifications ────────────────────────────────
#         if certifications:
#             add_section_heading('Certifications')
#             for cert in certifications:
#                 cert_line = g(cert.name)
#                 if g(cert.date): cert_line += f' | {g(cert.date)}'
#                 add_body(cert_line, bold=True)
#                 if g(cert.issuer):     add_meta(g(cert.issuer))
#                 if g(cert.verify_url): add_meta(f'Verify: {g(cert.verify_url)}')
#                 small_gap()
#
#         # ── 11. Achievements ──────────────────────────────────
#         if achievements:
#             add_section_heading('Honours & Achievements')
#             for ach in achievements:
#                 add_body(g(ach.title), bold=True)
#                 if g(ach.description):
#                     for line in g(ach.description).split('\n'):
#                         if line.strip():
#                             add_bullet(line)
#                 small_gap()
#
#         # ── 12. Projects ──────────────────────────────────────
#         if projects:
#             add_section_heading('Projects')
#             for proj in projects:
#                 add_body(g(proj.p_name), bold=True)
#
#                 # Project URL
#                 if g(proj.p_url):
#                     url = g(proj.p_url).strip()
#                     label = 'GitHub' if 'github' in url.lower() else 'Live'
#                     add_meta(f'{label}: {url}')
#
#                 # Description bullets from ProjectDetail
#                 try:
#                     from .models import ProjectDetail
#                     detail = ProjectDetail.objects.filter(project=proj).first()
#                     if detail and g(detail.description):
#                         for line in g(detail.description).split('\n'):
#                             if line.strip():
#                                 add_bullet(line)
#                 except Exception:
#                     pass
#
#                 small_gap()
#
#         # ── 13. Stream response ───────────────────────────────
#         buffer = io.BytesIO()
#         doc.save(buffer)
#         buffer.seek(0)
#
#         safe_name = full_name.replace(' ', '_').replace('/', '_')
#         filename = f'{safe_name}_CV.docx'
#
#         response = HttpResponse(
#             buffer,
#             content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
#         response['Content-Disposition'] = f'attachment; filename="{filename}"'
#         response['Cache-Control'] = 'no-cache'
#         return response
#
#     except Exception as e:
#         import traceback
#         traceback.print_exc()
#         return Response({'error': f'Word CV generation failed: {str(e)}'}, status=500)


# ─────────────────────────────────────────────────────────────
#  ADD THESE VIEWS TO YOUR views.py
# ─────────────────────────────────────────────────────────────

from .models import (
    ProfileLinks, ProfileSummary,
    EducationInfo, ExperienceInfo, CertificationInfo,
    AchievementInfo, LanguageSpoken, SoftSkillInfo, HobbyInfo,
)
from .serializer import (
    ProfileLinksSerializer, ProfileSummarySerializer,
    EducationInfoSerializer, ExperienceInfoSerializer,
    CertificationInfoSerializer, AchievementInfoSerializer,
    LanguageSpokenSerializer, SoftSkillInfoSerializer,
    HobbyInfoSerializer,
)


# ── Profile Links ─────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def profile_links(request):
    username = request.user.username
    obj, _   = ProfileLinks.objects.get_or_create(username=username)
    if request.method == 'GET':
        return Response(ProfileLinksSerializer(obj).data)
    serializer = ProfileLinksSerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Profile Summary ───────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def profile_summary(request):
    username = request.user.username
    obj, _   = ProfileSummary.objects.get_or_create(username=username)
    if request.method == 'GET':
        return Response(ProfileSummarySerializer(obj).data)
    serializer = ProfileSummarySerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Education ─────────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def education_list(request):
    username = request.user.username
    if request.method == 'GET':
        items = EducationInfo.objects.filter(username=username).order_by('order')
        return Response(EducationInfoSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = EducationInfoSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def education_detail(request, pk):
    try:
        obj = EducationInfo.objects.get(pk=pk, username=request.user.username)
    except EducationInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    if request.method == 'DELETE':
        obj.delete()
        return Response({'message': 'Deleted'})
    serializer = EducationInfoSerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Experience ────────────────────────────────────────────────

@api_view(['GET', 'POST'])
# @permission_classes([IsAuthenticated])
@permission_classes([AllowAny])
def experience_list(request):
    # username = request.user.username
    # if request.method == 'GET':
    if request.method == 'GET':
        # Public can view: Get username from the URL/Frontend
        username = request.query_params.get('username')
        if not username:
            return Response({"error": "Username is required"}, status=400)

        items = ExperienceInfo.objects.filter(username=username).order_by('order')
        return Response(ExperienceInfoSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = ExperienceInfoSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def experience_detail(request, pk):
    try:
        obj = ExperienceInfo.objects.get(pk=pk, username=request.user.username)
    except ExperienceInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    if request.method == 'DELETE':
        obj.delete()
        return Response({'message': 'Deleted'})
    serializer = ExperienceInfoSerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Certifications ────────────────────────────────────────────

# @api_view(['GET', 'POST'])
# # @permission_classes([IsAuthenticated])
# @permission_classes([AllowAny])
# def certification_list(request):
#     username = request.query_params.get('username')
#     # if request.method == 'GET':
#         # 1. Public can view: Get username from the URL/Frontend
#     if not username:
#         return Response({"error": "Username is required"}, status=400)
#             # username = request.user.username
#     # if request.method == 'GET':
#         items = CertificationInfo.objects.filter(username=username)
#         return Response(CertificationInfoSerializer(items, many=True).data)
#     data = {**request.data, 'username': username}
#     serializer = CertificationInfoSerializer(data=data)
#     if serializer.is_valid():
#         serializer.save()
#         return Response(serializer.data, status=201)
#     return Response(serializer.errors, status=400)
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from .models import CertificationInfo
from .serializer import CertificationInfoSerializer


@api_view(['GET', 'POST'])
@permission_classes([AllowAny])
def certification_list(request):
    # ── 1. PUBLIC GET REQUEST (Portfolio View) ──
    if request.method == 'GET':
        username = request.query_params.get('username')

        if not username:
            return Response({"error": "Username is required in URL parameters"}, status=400)

        items = CertificationInfo.objects.filter(username=username)
        return Response(CertificationInfoSerializer(items, many=True).data)

    # ── 2. PRIVATE POST REQUEST (Admin Dashboard) ──
    elif request.method == 'POST':
        # Manually check if they are logged in since the view allows anyone
        if not request.user.is_authenticated:
            return Response({"error": "Authentication required to add certifications"}, status=401)

        # Get the username from the secure token/session, NOT the URL
        username = request.user.username

        data = {**request.data, 'username': username}
        serializer = CertificationInfoSerializer(data=data)

        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=201)

        return Response(serializer.errors, status=400)



@api_view(['PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def certification_detail(request, pk):
    try:
        obj = CertificationInfo.objects.get(pk=pk, username=request.user.username)
    except CertificationInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    if request.method == 'DELETE':
        obj.delete()
        return Response({'message': 'Deleted'})
    serializer = CertificationInfoSerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Achievements ──────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def achievement_list(request):
    username = request.user.username
    if request.method == 'GET':
        items = AchievementInfo.objects.filter(username=username)
        return Response(AchievementInfoSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = AchievementInfoSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['PUT', 'DELETE'])
@permission_classes([IsAuthenticated])
def achievement_detail(request, pk):
    try:
        obj = AchievementInfo.objects.get(pk=pk, username=request.user.username)
    except AchievementInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    if request.method == 'DELETE':
        obj.delete()
        return Response({'message': 'Deleted'})
    serializer = AchievementInfoSerializer(obj, data=request.data, partial=True)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data)
    return Response(serializer.errors, status=400)


# ── Languages Spoken ──────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def language_list(request):
    username = request.user.username
    if request.method == 'GET':
        items = LanguageSpoken.objects.filter(username=username)
        return Response(LanguageSpokenSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = LanguageSpokenSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def language_detail(request, pk):
    try:
        obj = LanguageSpoken.objects.get(pk=pk, username=request.user.username)
    except LanguageSpoken.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    obj.delete()
    return Response({'message': 'Deleted'})


# ── Soft Skills ───────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def soft_skill_list(request):
    username = request.user.username
    if request.method == 'GET':
        items = SoftSkillInfo.objects.filter(username=username)
        return Response(SoftSkillInfoSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = SoftSkillInfoSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def soft_skill_detail(request, pk):
    try:
        obj = SoftSkillInfo.objects.get(pk=pk, username=request.user.username)
    except SoftSkillInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    obj.delete()
    return Response({'message': 'Deleted'})


# ── Hobbies ───────────────────────────────────────────────────

@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def hobby_list(request):
    username = request.user.username
    if request.method == 'GET':
        items = HobbyInfo.objects.filter(username=username)
        return Response(HobbyInfoSerializer(items, many=True).data)
    data = {**request.data, 'username': username}
    serializer = HobbyInfoSerializer(data=data)
    if serializer.is_valid():
        serializer.save()
        return Response(serializer.data, status=201)
    return Response(serializer.errors, status=400)


@api_view(['DELETE'])
@permission_classes([IsAuthenticated])
def hobby_detail(request, pk):
    try:
        obj = HobbyInfo.objects.get(pk=pk, username=request.user.username)
    except HobbyInfo.DoesNotExist:
        return Response({'error': 'Not found'}, status=404)
    obj.delete()
    return Response({'message': 'Deleted'})






class SettingsView(APIView):
    def get(self, request, username):
        obj = get_object_or_404(HomeInfo, username=username)
        serializer = SettingsSerializer(obj)
        return Response(serializer.data)

    def patch(self, request, username):
        obj = get_object_or_404(HomeInfo, username=username)
        serializer = SettingsSerializer(obj, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)





class CVPreferenceView(APIView):
    permission_classes = [AllowAny]

    def get(self, request, username):
        obj, _ = CVPreference.objects.get_or_create(username=username)
        serializer = CVPreferenceSerializer(obj)
        return Response(serializer.data)

    def patch(self, request, username):
        obj, _ = CVPreference.objects.get_or_create(username=username)
        serializer = CVPreferenceSerializer(obj, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)